import nltk
from docx import Document
import re

# NLTK veri dizinini ayarla
nltk.data.path.append('C:/Users/Ezat Kosif/Desktop/TextToSpech/nltk_data/')

# Punkt veri dosyasını kontrol et ve yükle
try:
    nltk.data.find('tokenizers/punkt')
    print("punkt data found")
except LookupError:
    print("punkt data not found. Please ensure it's correctly downloaded.")
    nltk.download('punkt')

def split_text_by_sentences(text):
    # Metni cümlelere ayır
    sentences = nltk.sent_tokenize(text, language='turkish')
    
    # Eğer cümle ayırıcı işaretler yoksa, metni belirli kurallar ve uzunluklarla böl
    if not sentences:
        # Örnek olarak metni belirli uzunluklarda parçalara ayırın
        sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
    
    return sentences

def read_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

def save_text_to_docx(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

# Word dosyasından metni oku
docx_file_path = "Contract Management.docx"
text = read_text_from_docx(docx_file_path)

# Cümlelere böl
sentences = split_text_by_sentences(text)

# Cümleleri tekrar birleştirerek metni oluştur
text_with_sentences = '\n'.join(sentences)

# Yeni metni aynı veya farklı bir Word belgesine kaydet
output_file_path = "New.docx"
save_text_to_docx(text_with_sentences, output_file_path)

print("Metin başarıyla kaydedildi.")
